# modules/reports/service/export_report_service.py

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from datetime import datetime
import os

from shared.dto.result import ServiceResult
from modules.reports.service.bao_cao_service import BaoCaoService


class ExportReportService:
    def __init__(self):
        self.svc = BaoCaoService()

    def export_word(self, loai: str, nam_hoc_id: int, ky: int, 
                    thang: int, export_path: str, options: dict):
        """Xuất báo cáo Word"""
        try:
            doc = Document()
            
            # Thiết lập font
            doc.styles['Normal'].font.name = 'Times New Roman'
            doc.styles['Normal'].font.size = Pt(12)
            
            # Tiêu đề
            self._add_header(doc, loai, nam_hoc_id, ky, thang)
            
            # Thống kê
            if options.get('thong_ke', True):
                self._add_statistics(doc, nam_hoc_id)
            
            # Bảng điểm chi tiết
            if options.get('chi_tiet', True):
                self._add_detail_table(doc, loai, nam_hoc_id, ky, thang)
            
            # Chữ ký
            if options.get('ky_ten', True):
                self._add_signature(doc)
            
            doc.save(export_path)
            return ServiceResult(ok=True, error="Xuất Word thành công")
        except Exception as e:
            return ServiceResult(ok=False, error=str(e))

    def export_pdf(self, loai: str, nam_hoc_id: int, ky: int, 
                   thang: int, export_path: str, options: dict):
        """Xuất báo cáo PDF"""
        try:
            doc = SimpleDocTemplate(export_path, pagesize=A4,
                                   topMargin=2*cm, bottomMargin=2*cm,
                                   leftMargin=2*cm, rightMargin=2*cm)
            
            story = []
            styles = getSampleStyleSheet()
            
            # Tạo style riêng
            title_style = ParagraphStyle(
                'TitleStyle', parent=styles['Title'],
                fontName='Helvetica-Bold', fontSize=16, alignment=TA_CENTER, spaceAfter=12
            )
            heading_style = ParagraphStyle(
                'HeadingStyle', parent=styles['Heading2'],
                fontName='Helvetica-Bold', fontSize=12, spaceAfter=6
            )
            normal_style = ParagraphStyle(
                'NormalStyle', parent=styles['Normal'],
                fontName='Helvetica', fontSize=10, alignment=TA_LEFT
            )
            
            # Tiêu đề
            self._add_pdf_header(story, title_style, loai, nam_hoc_id, ky, thang)
            
            # Thống kê
            if options.get('thong_ke', True):
                self._add_pdf_statistics(story, heading_style, normal_style, nam_hoc_id)
            
            # Bảng điểm
            if options.get('chi_tiet', True):
                self._add_pdf_table(story, heading_style, normal_style, loai, nam_hoc_id, ky, thang)
            
            # Chữ ký
            if options.get('ky_ten', True):
                self._add_pdf_signature(story, normal_style)
            
            doc.build(story)
            return ServiceResult(ok=True, error="Xuất PDF thành công")
        except Exception as e:
            return ServiceResult(ok=False, error=str(e))

    def _add_header(self, doc, loai, nam_hoc_id, ky, thang):
        """Thêm tiêu đề báo cáo"""
        # Tên trường
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("TRƯỜNG THCS PHONG BẮC")
        run.bold = True
        run.font.size = Pt(14)
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("SỞ GIÁO DỤC VÀ ĐÀO TẠO")
        run.bold = True
        run.font.size = Pt(14)
        
        # Đường kẻ
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("----------------------------------------")
        run.font.size = Pt(12)
        
        # Tiêu đề chính
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("BÁO CÁO TỔNG KẾT THI ĐUA")
        run.bold = True
        run.font.size = Pt(16)
        
        # Tên báo cáo
        loai_text = "GIÁO VIÊN" if loai == "giao_vien" else "HỌC SINH"
        
        # Xác định kỳ/tháng
        if ky == 1:
            ky_text = "HỌC KỲ I"
        elif ky == 2:
            ky_text = "HỌC KỲ II"
        else:
            ky_text = "CẢ NĂM"
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"NĂM HỌC 2026-2027 - {ky_text}")
        run.bold = True
        run.font.size = Pt(14)
        
        doc.add_paragraph()
        doc.add_paragraph()

    def _add_statistics(self, doc, nam_hoc_id):
        """Thêm phần thống kê"""
        p = doc.add_paragraph()
        run = p.add_run("I. THỐNG KÊ CHUNG")
        run.bold = True
        run.font.size = Pt(13)
        
        # Lấy dữ liệu thống kê
        result = self.svc.thong_ke_tong_quan(nam_hoc_id)
        if result.ok:
            data = result.data
            
            p = doc.add_paragraph()
            p.add_run(f"   - Tổng số giáo viên: {data.get('so_giao_vien', 0)}")
            
            p = doc.add_paragraph()
            p.add_run(f"   - Tổng số học sinh: {data.get('so_hoc_sinh', 0)}")
            
            p = doc.add_paragraph()
            p.add_run(f"   - Tổng số lớp: {data.get('so_lop', 0)}")
            
            p = doc.add_paragraph()
            p.add_run(f"   - Tổng số tổ chuyên môn: {data.get('so_to_chuyen_mon', 0)}")
        
        doc.add_paragraph()

    def _add_detail_table(self, doc, loai, nam_hoc_id, ky, thang):
        """Thêm bảng chi tiết"""
        p = doc.add_paragraph()
        run = p.add_run("II. BẢNG XẾP HẠNG")
        run.bold = True
        run.font.size = Pt(13)
        
        if loai == "giao_vien":
            result = self.svc.xep_hang_giao_vien(nam_hoc_id, ky, thang)
        else:
            result = self.svc.xep_hang_lop(nam_hoc_id, ky, thang)
        
        if result.ok:
            data = result.data
            
            if loai == "giao_vien":
                # Tạo bảng giáo viên
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Header
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "STT"
                hdr_cells[1].text = "Mã GV"
                hdr_cells[2].text = "Họ tên"
                hdr_cells[3].text = "Điểm"
                hdr_cells[4].text = "Xếp loại"
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Data
                for item in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get('thu_hang', 0))
                    row_cells[1].text = item.get('ma', '')
                    row_cells[2].text = item.get('ten', '')
                    row_cells[3].text = f"{item.get('diem', 0):.1f}"
                    row_cells[4].text = item.get('xep_loai', '')
                    
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                # Bảng lớp
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "STT"
                hdr_cells[1].text = "Mã lớp"
                hdr_cells[2].text = "Tên lớp"
                hdr_cells[3].text = "Điểm"
                hdr_cells[4].text = "Xếp thứ"
                
                for cell in hdr_cells:
                    cell.paragraphs[0].runs[0].bold = True
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                for item in data:
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(item.get('thu_hang', 0))
                    row_cells[1].text = item.get('ma', '')
                    row_cells[2].text = item.get('ten', '')
                    row_cells[3].text = f"{item.get('diem', 0):.3f}"
                    row_cells[4].text = str(item.get('thu_hang', 0))
                    
                    for cell in row_cells:
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()

    def _add_signature(self, doc):
        """Thêm chữ ký"""
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Ngày tháng
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run(f"Ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}")
        
        # Chữ ký
        for i in range(3):
            doc.add_paragraph()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("HIỆU TRƯỞNG")
        p.paragraph_format.space_before = Pt(30)

    # ===== PDF =====
    
    def _add_pdf_header(self, story, title_style, loai, nam_hoc_id, ky, thang):
        story.append(Paragraph("TRƯỜNG THCS PHONG BẮC", title_style))
        story.append(Paragraph("SỞ GIÁO DỤC VÀ ĐÀO TẠO", title_style))
        story.append(Spacer(1, 0.1*cm))
        story.append(Paragraph("BÁO CÁO TỔNG KẾT THI ĐUA", title_style))
        
        loai_text = "GIÁO VIÊN" if loai == "giao_vien" else "HỌC SINH"
        ky_text = "CẢ NĂM" if ky == 3 else ("HỌC KỲ I" if ky == 1 else "HỌC KỲ II")
        story.append(Paragraph(f"NĂM HỌC 2026-2027 - {ky_text}", title_style))
        story.append(Spacer(1, 0.5*cm))

    def _add_pdf_statistics(self, story, heading_style, normal_style, nam_hoc_id):
        story.append(Paragraph("I. THỐNG KÊ CHUNG", heading_style))
        
        result = self.svc.thong_ke_tong_quan(nam_hoc_id)
        if result.ok:
            data = result.data
            story.append(Paragraph(f"   - Tổng số giáo viên: {data.get('so_giao_vien', 0)}", normal_style))
            story.append(Paragraph(f"   - Tổng số học sinh: {data.get('so_hoc_sinh', 0)}", normal_style))
            story.append(Paragraph(f"   - Tổng số lớp: {data.get('so_lop', 0)}", normal_style))
            story.append(Paragraph(f"   - Tổng số tổ chuyên môn: {data.get('so_to_chuyen_mon', 0)}", normal_style))
        
        story.append(Spacer(1, 0.3*cm))

    def _add_pdf_table(self, story, heading_style, normal_style, loai, nam_hoc_id, ky, thang):
        story.append(Paragraph("II. BẢNG XẾP HẠNG", heading_style))
        
        if loai == "giao_vien":
            result = self.svc.xep_hang_giao_vien(nam_hoc_id, ky, thang)
        else:
            result = self.svc.xep_hang_lop(nam_hoc_id, ky, thang)
        
        if result.ok:
            data = result.data[:20]
            
            if loai == "giao_vien":
                table_data = [["STT", "Mã GV", "Họ tên", "Điểm", "Xếp loại"]]
                for item in data:
                    table_data.append([
                        str(item.get('thu_hang', 0)),
                        item.get('ma', ''),
                        item.get('ten', ''),
                        f"{item.get('diem', 0):.3f}",
                        item.get('xep_loai', '')
                    ])
            else:
                table_data = [["STT", "Mã lớp", "Tên lớp", "Điểm", "Xếp thứ"]]
                for item in data:
                    table_data.append([
                        str(item.get('thu_hang', 0)),
                        item.get('ma', ''),
                        item.get('ten', ''),
                        f"{item.get('diem', 0):.3f}",
                        str(item.get('thu_hang', 0))
                    ])
            
            # Tạo bảng PDF
            table = Table(table_data, colWidths=[1*cm, 2*cm, 4*cm, 2*cm, 3*cm])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1D9E75')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            story.append(table)
        
        story.append(Spacer(1, 0.3*cm))

    def _add_pdf_signature(self, story, normal_style):
        story.append(Spacer(1, 1*cm))
        
        # Ngày tháng
        story.append(Paragraph(f"Ngày {datetime.now().day} tháng {datetime.now().month} năm {datetime.now().year}", 
                              normal_style))
        
        story.append(Spacer(1, 2*cm))
        
        # Chữ ký
        story.append(Paragraph("HIỆU TRƯỞNG", normal_style))